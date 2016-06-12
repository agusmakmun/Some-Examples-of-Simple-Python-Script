import os, uuid
from docx import Document
from docx.shared import Inches

docx_file   = 'penelitian.docx'
path_images = 'images'

def insetImgToDocx(image):
    document = Document(docx_file)
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(image, width=Inches(5.5)) #inces sesuai yang ada di garis
    r.add_text('Ini picture: {}'.format(image))
    document.save(docx_file)

def takeScreenshot():
    generateID = uuid.uuid4().hex + ".png"
    os.system("import -window root {}/{}".format(path_images, generateID))
    insetImgToDocx(path_images+"/"+generateID)

if __name__ == '__main__':
    takeScreenshot()
